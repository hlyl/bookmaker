#!/usr/bin/env python
import warnings

# Suppress specific warnings
warnings.filterwarnings("ignore", category=UserWarning, module="ebooklib.epub")
warnings.filterwarnings("ignore", category=FutureWarning, module="ebooklib.epub")

import os
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from typing import Dict, List, Tuple

from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import argparse

# Load environment variables from .env file
load_dotenv()
# Initialize the OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Function to summarize text with specific instructions
def summarize_text_with_chatgpt(text):
    chat_completion = client.chat.completions.create(
        model="gpt-4o-2024-05-13",
        messages=[
            {
                "role": "system",
                "content": (
                "Summarize the key takeaways from the provided text into a 250-word section. "
                "Then, condense the main points from the entire text into a recap that is "
                "approximately 75% of the original word count. Please write the summary and "
                "recap in a book-like format, using clear and concise language to reflect "
                "on the learnings from the text. Ensure the text is formatted for optimal "
                "readability, using headings, paragraphs, and proper spacing to make it "
                "easy to follow and understand"
                ),
            },
            {"role": "user", "content": text},
        ],
    )
    return chat_completion.choices[0].message.content.strip()

def read_epub_book(file_name: str) -> epub.EpubBook:
    """Reads an EPUB book from the given file name."""
    return epub.read_epub(file_name)

def list_epub_items(book: epub.EpubBook):
    """Lists all items in the EPUB book for debugging purposes."""
    for item in book.get_items():
        if isinstance(item, epub.EpubHtml):
            print(f"ID: {item.get_id()}, Type: {item.get_type()}, Name: {item.get_name()}")
        else:
            print(f"ID: {item.get_id()}, Type: {item.get_type()}")

def extract_chapters_from_ncx(book: epub.EpubBook) -> List[Tuple[str, str]]:
    """Extracts chapters from an EPUB book using the NCX file."""
    chapters = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_NAVIGATION:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            for navpoint in soup.find_all("navpoint"):
                title = navpoint.find("navlabel").get_text(strip=True)
                src = navpoint.find("content")["src"]
                chapters.append((title, src.split('#')[0]))
    return chapters

def chapter_to_str(chapter: epub.EpubHtml) -> str:
    """Converts a chapter to a string of text."""
    soup = BeautifulSoup(chapter.get_body_content(), "html.parser")
    text = [para.get_text() for para in soup.find_all("p")]
    return " ".join(text)

def extract_chapters_to_text(book: epub.EpubBook, test_mode: bool = False) -> Dict[str, str]:
    """Extracts chapters from an EPUB book and converts them to text with their titles."""
    chapters = extract_chapters_from_ncx(book)
    texts = {}
    for i, (title, src) in enumerate(chapters):
        if test_mode and i >= 3:
            break
        chapter = book.get_item_with_href(src)
        if chapter is not None:
            chapter_text = chapter_to_str(chapter)
            if chapter_text.strip():  # Ensure the chapter text is not empty
                texts[title] = chapter_text
            else:
                print(f"Warning: Chapter '{title}' is empty.")
        else:
            print(f"Warning: Chapter with href {src} not found.")
    return texts

def count_words(text: str) -> int:
    """Counts the number of words in a given text."""
    words = text.split()
    return len(words)

def main():
    parser = argparse.ArgumentParser(description="Process an EPUB file and generate a summary DOCX file.")
    parser.add_argument('file_name', type=str, help='The path to the EPUB file to be processed')
    parser.add_argument('--test', action='store_true', help='Limit processing to the first 3 chapters')
    args = parser.parse_args()

    file_name = args.file_name
    test_mode = args.test
    book = read_epub_book(file_name)

    # List all items in the EPUB for debugging
    list_epub_items(book)

    chapters_texts = extract_chapters_to_text(book, test_mode)

    if not chapters_texts:
        print("No chapters found or extracted.")
        return

    # Create a Word document to save the summaries
    document = Document()

    for chapter_name, chapter_text in chapters_texts.items():
        print(f"Processing chapter: {chapter_name}...")  # Debug log
        summarized_text = summarize_text_with_chatgpt(chapter_text)
        print(f"Chapter: {chapter_name}")
        print(f"Original Text: {chapter_text[:500]}...")  # Debug log, first 500 chars
        print(f"Summarized Text: {summarized_text}")

        # Add chapter name
        document.add_heading(chapter_name, level=1)

        # Split the summarized text by sections to avoid repeated headings
        sections = summarized_text.split('\n')
        added_sections = set()
        for section in sections:
            if section.startswith("Key Takeaway(s):") and "Key Takeaway(s):" not in added_sections:
                document.add_paragraph(section)
                added_sections.add("Key Takeaway(s):")
            elif section.startswith("Recap:") and "Recap:" not in added_sections:
                document.add_paragraph(section)
                added_sections.add("Recap:")
            else:
                document.add_paragraph(section)

        # Add a section break (next page)
        document.add_page_break()

    # Determine the output file name based on the input file name
    base_name = os.path.splitext(os.path.basename(file_name))[0]
    output_file = f"{base_name}_Summary.docx"
    document.save(output_file)
    print(f"Summarized chapters saved to {output_file}")

if __name__ == "__main__":
    main()
